"""
Syllabus Parser for SCDO
Extracts structured information from syllabus documents (PDF, DOCX, TXT)
"""

import re
from typing import Dict, List, Optional, Any
from pathlib import Path
import logging

try:
    import pypdf
    import pdfplumber
except ImportError:
    pypdf = None
    pdfplumber = None
    logging.warning("PDF libraries not installed. Install with: pip install pypdf pdfplumber")

try:
    from docx import Document
except ImportError:
    logging.warning("python-docx not installed. Install with: pip install python-docx")

# OCR support for scanned PDFs
try:
    import pytesseract
    from pdf2image import convert_from_path
    from PIL import Image
    import os
    import platform
    
    # Configure Tesseract path for Windows
    if platform.system() == 'Windows':
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
            os.path.expanduser(r'~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'),
        ]
        for tess_path in tesseract_paths:
            if os.path.exists(tess_path):
                pytesseract.pytesseract.tesseract_cmd = tess_path
                logging.info(f"Tesseract found at: {tess_path}")
                break
    
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logging.info("OCR libraries not installed. Scanned PDFs won't be supported. Install with: pip install pytesseract pdf2image Pillow")

from ..utils.text_processing import TextProcessor

# Import AI model for LLM-based parsing fallback
try:
    from ..ai.openrouter_model import OpenRouterModel
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    logging.info("AI model not available for LLM-based parsing fallback")


class SyllabusParser:
    """Parse syllabus documents and extract structured information"""
    
    def __init__(self, use_llm_fallback: bool = True):
        self.logger = logging.getLogger(__name__)
        self.text_processor = TextProcessor()
        self.use_llm_fallback = use_llm_fallback and LLM_AVAILABLE
        self.ai_model = None
        
        if self.use_llm_fallback:
            try:
                self.ai_model = OpenRouterModel()
                if not self.ai_model.is_available():
                    self.logger.info("LLM not available (no API key), using regex-only parsing")
                    self.use_llm_fallback = False
            except Exception as e:
                self.logger.warning(f"Failed to initialize AI model for parsing: {e}")
                self.use_llm_fallback = False
        
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse syllabus file and extract structured data
        
        Args:
            file_path: Path to syllabus file
            
        Returns:
            Structured syllabus data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        # Determine file type and parse
        if file_path.suffix.lower() == '.pdf':
            text = self._parse_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            text = self._parse_docx(file_path)
        elif file_path.suffix.lower() == '.txt':
            text = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
            
        # Extract structured information
        return self._extract_structure(text)
        
    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF, using table-aware extraction when available"""
        text = ""

        if pdfplumber is None and pypdf is None:
            raise ImportError("PDF libraries not installed. Please install pypdf and pdfplumber.")

        self.logger.info(f"Parsing PDF: {file_path}")

        try:
            # Try pdfplumber first (better for tables)
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    self.logger.info(f"PDF has {len(pdf.pages)} pages")
                    for i, page in enumerate(pdf.pages):
                        page_text = self._extract_page_text(page)
                        if page_text:
                            text += page_text + "\n"
                            self.logger.debug(f"Page {i+1}: extracted {len(page_text)} chars")
                        else:
                            self.logger.warning(f"Page {i+1}: no text extracted")
                self.logger.info(f"pdfplumber extracted total {len(text)} chars")
            else:
                raise ImportError("pdfplumber not installed")

        except Exception as e:
            self.logger.warning(f"pdfplumber failed or missing, trying pypdf: {e}")

            # Fallback to pypdf
            if pypdf:
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = pypdf.PdfReader(f)
                        self.logger.info(f"pypdf: PDF has {len(pdf_reader.pages)} pages")
                        for i, page in enumerate(pdf_reader.pages):
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                                self.logger.debug(f"Page {i+1}: extracted {len(page_text)} chars")
                    self.logger.info(f"pypdf extracted total {len(text)} chars")
                except Exception as e2:
                    self.logger.error(f"PDF parsing failed: {e2}")
                    raise
            else:
                if not text: # Only raise if we haven't extracted anything yet and both failed/missing
                     raise ImportError("pypdf not installed and pdfplumber failed") from e

        # If no text extracted or very minimal text (likely garbage from scanned PDF), try OCR
        # Using 50 chars as threshold since scanned PDFs often extract some whitespace/garbage
        extracted_text_length = len(text.strip())
        if extracted_text_length < 50:
            self.logger.info(f"Minimal text extracted ({extracted_text_length} chars), attempting OCR for scanned PDF...")
            if OCR_AVAILABLE:
                self.logger.info("Using local Tesseract OCR...")
                ocr_text = self._ocr_pdf(file_path)
            else:
                # Try online OCR if local OCR is not available
                self.logger.info("Local OCR not available, attempting online OCR (OCR.space API)...")
                ocr_text = self._ocr_online(file_path)
            
            # Use OCR text if it has more content than original extraction
            if len(ocr_text.strip()) > extracted_text_length:
                self.logger.info(f"OCR extracted {len(ocr_text)} chars (more than original {extracted_text_length})")
                text = ocr_text
            else:
                self.logger.warning(f"OCR did not improve extraction (got {len(ocr_text.strip())} chars)")
            
        if len(text.strip()) > 0:
            self.logger.info(f"PDF extraction successful: {len(text.strip())} total characters")
        else:
            self.logger.error("Failed to extract text from PDF (likely a scanned document)")
            if not OCR_AVAILABLE:
                self.logger.error("Install OCR to handle scanned PDFs: pip install pytesseract pdf2image Pillow")
                self.logger.error("Also install Tesseract: https://github.com/tesseract-ocr/tesseract")
                
        return text

    def _extract_page_text(self, page) -> str:
        """
        Extract text from a single pdfplumber page.

        Uses table-aware extraction: if the page contains tables, reconstructs
        text from table cells in proper reading order. Falls back to plain
        extract_text() for pages without tables.

        Reading order: interleaves non-table text and tables based on their
        vertical position on the page, so headings above a table appear before
        the table, and text below appears after.
        """
        # Try table-aware extraction first
        try:
            tables = page.find_tables()
        except Exception:
            tables = []

        if not tables:
            return page.extract_text() or ""

        table_bboxes = [t.bbox for t in tables]  # (x0, top, x1, bottom)

        # Collect non-table chars and group into lines
        chars = page.chars
        non_table_chars = [
            c for c in chars
            if not any(
                bbox[0] <= c['x0'] <= bbox[2] and bbox[1] <= c['top'] <= bbox[3]
                for bbox in table_bboxes
            )
        ]

        # Group non-table chars into text lines by vertical position
        non_table_lines = []
        if non_table_chars:
            # Tolerance for same-line grouping (chars within 3px vertically)
            sorted_chars = sorted(non_table_chars, key=lambda c: (c['top'], c['x0']))
            current_line = [sorted_chars[0]]
            for c in sorted_chars[1:]:
                if abs(c['top'] - current_line[0]['top']) <= 3:
                    current_line.append(c)
                else:
                    line_text = "".join(ch['text'] for ch in sorted(current_line, key=lambda x: x['x0']))
                    if line_text.strip():
                        non_table_lines.append((current_line[0]['top'], line_text))
                    current_line = [c]
            # Flush last line
            line_text = "".join(ch['text'] for ch in sorted(current_line, key=lambda x: x['x0']))
            if line_text.strip():
                non_table_lines.append((current_line[0]['top'], line_text))

        # Collect table text with their vertical position
        table_entries = []
        for table, bbox in zip(tables, table_bboxes):
            rows = table.extract()
            if not rows:
                continue
            row_texts = []
            for row in rows:
                cells = [(c or "") for c in row]
                line = "\t".join(cells)
                if line.strip():
                    row_texts.append(line)
            if row_texts:
                table_entries.append((bbox[1], "\n".join(row_texts)))  # bbox[1] = top

        # Merge both lists and sort by vertical position (reading order)
        all_blocks = non_table_lines + table_entries
        all_blocks.sort(key=lambda b: b[0])

        return "\n".join(text for _, text in all_blocks)
    
    def _ocr_online(self, file_path: Path) -> str:
        """Extract text from scanned PDF using free online OCR API (OCR.space)"""
        import requests
        import os
        
        text = ""
        
        # Check file size - free API has 1MB limit
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > 1:
            self.logger.warning(f"File is {file_size_mb:.1f}MB - OCR.space free API limit is 1MB")
            self.logger.info("Consider using a smaller PDF or installing local Tesseract OCR")
        
        # Retry logic for network issues
        max_retries = 2
        for attempt in range(max_retries):
            try:
                self.logger.info(f"OCR.space API attempt {attempt + 1}/{max_retries}...")
                
                # OCR.space free API - no signup required for basic usage
                api_url = "https://api.ocr.space/parse/image"
                
                with open(file_path, 'rb') as f:
                    response = requests.post(
                        api_url,
                        files={'file': f},
                        data={
                            'apikey': 'helloworld',  # Free demo API key
                            'language': 'eng',
                            'isOverlayRequired': False,
                            'filetype': 'PDF',
                            'detectOrientation': True,
                            'scale': True,
                            'OCREngine': 2  # Better for dense text
                        },
                        timeout=120  # Increased timeout for large files
                    )
                
                if response.status_code == 200:
                    result = response.json()
                    if result.get('ParsedResults'):
                        for page_result in result['ParsedResults']:
                            page_text = page_result.get('ParsedText', '')
                            if page_text:
                                text += page_text + "\n"
                        if text:
                            self.logger.info(f"Online OCR extracted {len(text)} chars")
                            return text  # Success, return immediately
                    else:
                        error_msg = result.get('ErrorMessage', '') or result.get('OCRExitCode', 'Unknown error')
                        self.logger.warning(f"OCR.space returned no results: {error_msg}")
                        if 'limit' in str(error_msg).lower() or 'size' in str(error_msg).lower():
                            self.logger.info("File may be too large for free API")
                            break  # Don't retry if it's a size limit issue
                else:
                    self.logger.error(f"OCR.space API error: {response.status_code}")
                    
            except requests.exceptions.Timeout:
                self.logger.warning(f"OCR timeout on attempt {attempt + 1}")
                if attempt < max_retries - 1:
                    self.logger.info("Retrying...")
                    continue
            except Exception as e:
                self.logger.error(f"Online OCR failed: {e}")
                break
            
        return text
    
    def _ocr_pdf(self, file_path: Path) -> str:
        """Extract text from scanned PDF using local OCR (Tesseract)"""
        text = ""
        try:
            self.logger.info(f"Converting PDF pages to images for OCR...")
            # Convert PDF pages to images at 300 DPI for better OCR quality
            # Higher DPI = more text extracted (tested: 300 DPI extracts 3x more than 200 DPI)
            images = convert_from_path(str(file_path), dpi=300)
            self.logger.info(f"Converted {len(images)} pages to images")
            
            for i, image in enumerate(images):
                self.logger.info(f"OCR processing page {i+1}...")
                page_text = pytesseract.image_to_string(image, lang='eng')
                if page_text:
                    text += page_text + "\n"
                    self.logger.info(f"Page {i+1}: OCR extracted {len(page_text)} chars")
                else:
                    self.logger.warning(f"Page {i+1}: OCR found no text")
                    
            self.logger.info(f"OCR extracted total {len(text)} chars")
            
        except Exception as e:
            self.logger.error(f"Local OCR failed: {e}")
            self.logger.info("Make sure Tesseract OCR is installed: https://github.com/tesseract-ocr/tesseract")
            
        return text
        
    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join([cell.text for cell in row.cells])
                    
            return text
            
        except Exception as e:
            self.logger.error(f"DOCX parsing failed: {e}")
            raise
            
    def _parse_txt(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
                
    def _extract_first_course_text(self, text: str) -> str:
        """
        Extract only the first course's text from a multi-course PDF.
        
        Detects course boundaries by looking for course code patterns and
        returns only the text for the first complete course.
        
        Args:
            text: Raw PDF text potentially containing multiple courses
            
        Returns:
            Text for the first course only
        """
        # Pattern to detect course headers like "MSCCS24101: Design and Analysis"
        # or "Course Code Course Name" table format
        course_header_patterns = [
            # Pattern: "COURSECODE: Course Title" followed by Course Type
            r'([A-Z]{2,}\d{5,}):\s*[A-Za-z][^\n]+\nCourse\s*Type',
            # Pattern: Course code in structured section
            r'([A-Z]{2,}\d{5,})\s+[A-Za-z][A-Za-z\s]+\nCourse\s*Type',
        ]
        
        # Find all course code occurrences that mark course starts
        course_starts = []
        for pattern in course_header_patterns:
            for match in re.finditer(pattern, text):
                course_starts.append(match.start())
        
        # Sort and remove duplicates
        course_starts = sorted(set(course_starts))
        
        if len(course_starts) >= 2:
            # Multi-course document - extract only first course
            first_course_start = course_starts[0]
            second_course_start = course_starts[1]
            
            # Include some context before the first course (headers/metadata)
            prefix_start = max(0, first_course_start - 500)
            
            first_course_text = text[prefix_start:second_course_start]
            self.logger.info(f"Multi-course PDF detected. Extracting first course only (chars {first_course_start}-{second_course_start})")
            return first_course_text
        
        # Alternative detection: look for repeated "Course Content" sections
        course_content_matches = list(re.finditer(r'Course\s*Content\nUnit\s*No', text, re.IGNORECASE))
        
        if len(course_content_matches) >= 2:
            # Found multiple course content sections
            first_end = course_content_matches[1].start()
            
            # Find where the first course starts (after Programme Structure)
            prog_struct_match = re.search(r'Programme\s*Structure', text)
            first_start = prog_struct_match.end() if prog_struct_match else 0
            
            first_course_text = text[first_start:first_end]
            self.logger.info(f"Detected {len(course_content_matches)} courses. Extracting first course only.")
            return first_course_text
        
        # Single course or couldn't detect boundaries - use full text
        return text

    def _extract_structure(self, text: str) -> Dict[str, Any]:
        """
        Extract structured information from syllabus text
        
        Args:
            text: Raw syllabus text
            
        Returns:
            Structured syllabus data
        """
        self.logger.info(f"Extracting structure from text of {len(text)} chars")
        
        # Extract only the first course's text if this is a multi-course document
        course_text = self._extract_first_course_text(text)
        
        # First, try regex-based extraction
        structure = {
            'course_title': self._extract_course_title(course_text),
            'course_code': self._extract_course_code(course_text),
            'credits': self._extract_credits(course_text),
            'prerequisites': self._extract_prerequisites(course_text),
            'objectives': self._extract_objectives(course_text),
            'learning_outcomes': self._extract_learning_outcomes(course_text),
            'units': self._extract_units(course_text),
            'assessment_pattern': self._extract_assessment(course_text),
            'references': self._extract_references(course_text),
            'co_po_mapping': self._extract_co_po_mapping(course_text),
            'raw_text': text  # Keep original full text for reference
        }
        
        # Check if regex extraction yielded poor results - use LLM fallback
        regex_units = len(structure['units'])
        regex_outcomes = len(structure['learning_outcomes'])
        
        if self.use_llm_fallback and self.ai_model and (regex_units < 2 or regex_outcomes < 2):
            self.logger.info(f"Regex extraction yielded poor results (units={regex_units}, outcomes={regex_outcomes}). Trying LLM fallback...")
            
            try:
                llm_structure = self._extract_structure_with_llm(course_text)
                
                # Merge LLM results - prefer LLM if it found more content
                if len(llm_structure.get('units', [])) > regex_units:
                    self.logger.info(f"LLM found {len(llm_structure['units'])} units vs regex's {regex_units}")
                    structure['units'] = llm_structure['units']
                    
                if len(llm_structure.get('learning_outcomes', [])) > regex_outcomes:
                    self.logger.info(f"LLM found {len(llm_structure['learning_outcomes'])} outcomes vs regex's {regex_outcomes}")
                    structure['learning_outcomes'] = llm_structure['learning_outcomes']
                
                # Update other fields if regex got defaults
                if structure['course_title'] == 'Unknown Course' and llm_structure.get('course_title'):
                    structure['course_title'] = llm_structure['course_title']
                if structure['course_code'] == 'UNKNOWN' and llm_structure.get('course_code'):
                    structure['course_code'] = llm_structure['course_code']
                    
            except Exception as e:
                self.logger.warning(f"LLM fallback parsing failed: {e}")
        
        self.logger.info(f"Extraction complete: title='{structure['course_title']}', outcomes={len(structure['learning_outcomes'])}, units={len(structure['units'])}, raw_text={len(structure['raw_text'])} chars")
        
        return structure
    
    def _extract_structure_with_llm(self, text: str) -> Dict[str, Any]:
        """
        Use LLM to extract structured syllabus information.
        More robust for OCR'd text with errors or unusual formatting.
        """
        import json
        
        prompt = f"""Extract structured information from this syllabus text. The text may contain OCR errors.

Return a JSON object with these fields:
- course_title: string (the course name)
- course_code: string (e.g., "BTECCE25101")
- units: array of objects with:
  - unit_number: string or number
  - title: string (unit title)
  - topics: array of topic strings
  - hours: number (teaching hours)
- learning_outcomes: array of objects with:
  - code: string (e.g., "CO1", "CO2")
  - description: string (the learning outcome text)
  - bloom_level: string (one of: "Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create")

Syllabus text:
{text[:6000]}

Return ONLY valid JSON, no markdown formatting or code blocks."""

        system_prompt = """You are a syllabus parser assistant. Extract structured information from educational syllabus documents.
Be robust to OCR errors (e.g., 'Howrs' means 'Hours', '0' might be 'O', etc.).
Focus on finding units and learning outcomes even if the text is messy."""

        try:
            response = self.ai_model.generate(
                prompt=prompt,
                system_prompt=system_prompt,
                temperature=0.1,  # Low temp for structured extraction
                max_tokens=2000
            )
            
            # Clean and parse JSON
            response = response.strip()
            if response.startswith('```'):
                # Remove markdown code blocks if present
                response = re.sub(r'^```(?:json)?\n?', '', response)
                response = re.sub(r'\n?```$', '', response)
            
            result = json.loads(response)
            
            self.logger.info(f"LLM extracted: {len(result.get('units', []))} units, {len(result.get('learning_outcomes', []))} outcomes")
            return result
            
        except json.JSONDecodeError as e:
            self.logger.warning(f"LLM returned invalid JSON: {e}")
            return {}
        except Exception as e:
            self.logger.warning(f"LLM extraction failed: {e}")
            return {}
        
    def _extract_course_title(self, text: str) -> str:
        """Extract course title"""
        patterns = [
            # Pattern for "COURSECODE: Course Title" format
            r'([A-Z]{2,}[\d]+):\s*(.+?)(?=\n)',
            # Pattern for "Course Title:" format
            r'(?:Course Title|Course Name)[:\s]+(.+?)(?=\n|Course Code)',
            # Pattern for title followed by Course Type
            r'([A-Za-z][A-Za-z\s&]+?)\nCourse\s*Type',
            # Subject pattern
            r'(?:Subject|Course)[:\s]+(.+?)(?=\n)',
            # Pattern for "Title: Something" at start
            r'^(.+?)\n(?:Course\s*Type|Credits?|L-T-P)',
            # Pattern for syllabus headers
            r'Syllabus\s*(?:for|of)?\s*:?\s*(.+?)(?=\n)',
            # Pattern for "Name of the Course:" 
            r'(?:Name of the course|Course name)[:\s]+(.+?)(?=\n)',
            # Pattern for lines after department headers
            r'(?:Department\s*of\s*.+?\n+)(.+?)(?=\n)',
            # Pattern matching any bolded or emphasized title
            r'\*\*(.+?)\*\*',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                # Handle patterns with multiple groups
                if len(match.groups()) >= 2 and match.group(2):
                    title = match.group(2).strip()
                else:
                    title = match.group(1).strip()
                # Clean up title - remove extra whitespace and common prefixes
                title = re.sub(r'\s+', ' ', title)
                title = re.sub(r'^(Course|Subject|Title)[:\s]*', '', title, flags=re.IGNORECASE)
                if title and len(title) > 3:
                    return title
        
        # Fallback: Try to find first meaningful line (often the title)
        lines = text.split('\n')
        for line in lines[:15]:  # Check first 15 lines
            line = line.strip()
            # Look for lines that look like titles (mixed case, reasonable length)
            if line and len(line) > 10 and len(line) < 100:
                if not re.match(r'^(Course\s*Type|Credits?|L-T-P|Semester|University|Department|Programme)', line, re.IGNORECASE):
                    if re.match(r'[A-Z]', line):  # Starts with capital letter
                        return line
                
        return "Unknown Course"
        
    def _extract_course_code(self, text: str) -> str:
        """Extract course code"""
        patterns = [
            # Pattern for course codes like MSCCS24101, CSE301, CS101, etc.
            r'(?:Course\s*Code|Code)[:\s]+([A-Z]{2,}[\d]+[A-Z]?\d*)',
            r'\b([A-Z]{2,}\d{4,6})\b',
            r'\b([A-Z]{2,4}\d{3,4})\b',
            # Pattern for codes at the start of title line
            r'^([A-Z]{2,}\d+)[:\s]',
            # Pattern for codes in table format
            r'Course\s*Code\s*\n?\s*([A-Z]{2,}[\d]+)',
            # Pattern for codes like 22CS301
            r'\b(\d{2}[A-Z]{2,}\d{3})\b',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
            if match:
                code = match.group(1).upper()
                # Validate - should have both letters and numbers
                if re.match(r'^[A-Z]+\d+$', code) or re.match(r'^\d+[A-Z]+\d+$', code):
                    # Exclude CO patterns (CO1, CO2, etc.) - these are course outcomes, not course codes
                    if re.match(r'^CO\d+$', code):
                        continue
                    return code
        
        # Fallback: Look for any alphanumeric pattern that looks like a code (at least 5 chars)
        code_match = re.search(r'\b([A-Z]{2,}\d{3,}[A-Z]?\d*)\b', text)
        if code_match:
            code = code_match.group(1).upper()
            if not re.match(r'^CO\d+$', code):
                return code
            
        return "UNKNOWN"
        
    def _extract_credits(self, text: str) -> str:
        """Extract credit hours (L-T-P format)"""
        patterns = [
            # Standard L-T-P format
            r'(?:L-T-P|L T P|LTP)[:\s]*([\d]+)\s*[-:]\s*([\d]+)\s*[-:]\s*([\d]+)',
            r'(?:Credits?|Credit Hours?)[:\s]+([\d]+[-][\d]+[-][\d]+)',
            r'([\d]+)\s*[-]\s*([\d]+)\s*[-]\s*([\d]+)\s*(?:Credits?|Hours?)?',
            # Table format: L  T  P  Total
            r'L\s+T\s+P\s+Total\s*\n?\s*([\d]+)\s+([\d]+)\s+([\d]+)',
            # Credits as single number
            r'(?:Credits?|Credit\s*Hours?)[:\s]+(\d+)\s*$',
            # Pattern for "3:1:0" or "3 : 1 : 0"
            r'(\d+)\s*:\s*(\d+)\s*:\s*(\d+)',
            # Teaching Hours format
            r'(?:Teaching\s*Hours?|Lecture\s*Hours?)[:\s]*(\d+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                if len(match.groups()) == 3:
                    return f"{match.group(1)}-{match.group(2)}-{match.group(3)}"
                elif len(match.groups()) == 1:
                    # Single credit value - assume it's lecture only
                    credit = match.group(1)
                    if '-' in credit:
                        return credit
                    return f"{credit}-0-0"
        
        # Fallback: Look for any X-X-X pattern in first 500 chars
        ltp_match = re.search(r'\b(\d)\s*[-:]\s*(\d)\s*[-:]\s*(\d)\b', text[:500])
        if ltp_match:
            return f"{ltp_match.group(1)}-{ltp_match.group(2)}-{ltp_match.group(3)}"
                
        return "3-0-0"  # Default to typical lecture course
        
    def _extract_prerequisites(self, text: str) -> List[str]:
        """Extract prerequisites"""
        pattern = r'(?:Prerequisites?|Pre-requisites?)[:\s]+(.+?)(?=\n\n|\n[A-Z])'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            prereq_text = match.group(1)
            # Split by common delimiters
            prereqs = re.split(r'[,;]|\band\b', prereq_text)
            return [p.strip() for p in prereqs if p.strip()]
            
        return []
        
    def _extract_objectives(self, text: str) -> List[str]:
        """Extract course objectives"""
        pattern = r'(?:Course Objectives?|Objectives?)[:\s]+(.+?)(?=\n\n|Course Outcomes?|Learning Outcomes?)'
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            obj_text = match.group(1)
            # Extract numbered or bulleted items
            objectives = re.findall(r'(?:\d+\.|•|\*)\s*(.+?)(?=\n\d+\.|\n•|\n\*|$)', obj_text, re.DOTALL)
            return [obj.strip() for obj in objectives if obj.strip()]
            
        return []
        
    def _extract_learning_outcomes(self, text: str) -> List[Dict[str, str]]:
        """Extract course outcomes with Bloom's classification"""
        outcomes = []
        
        # First, try to find the Course Outcomes section
        co_section_pattern = r'Course\s*Outcomes?(.+?)(?:Mapping|CO\s*PO|Unit|Course\s*Content|$)'
        co_section_match = re.search(co_section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if co_section_match:
            co_text = co_section_match.group(1)
            
            # Pattern for numbered outcomes: "1 Analyze worst-case..."
            numbered_pattern = r'^\s*(\d+)\s+([A-Z][^\n]+?)(?=\n\s*\d+\s+[A-Z]|\nMapping|$)'
            matches = list(re.finditer(numbered_pattern, co_text, re.MULTILINE | re.DOTALL))
            
            for match in matches:
                num = match.group(1)
                description = match.group(2).strip()
                
                # Skip if it looks like a CO-PO mapping row (contains mostly numbers)
                if re.match(r'^[\d\s,.-]+$', description):
                    continue
                    
                # Skip very short descriptions
                if len(description) < 20:
                    continue
                
                bloom_level = self.text_processor.classify_bloom_level(description)
                
                outcomes.append({
                    'code': f'CO{num}',
                    'description': description,
                    'bloom_level': bloom_level
                })
            
            # NEW: Handle flattened PDF table format where CO numbers and statements are on separate lines
            # Format: "CO No.\n1\n2\n3...\nStatement\nFirst outcome...\nSecond outcome..."
            if not outcomes:
                # Look for "Statement" header followed by outcomes
                statement_match = re.search(r'Statement\n(.+?)(?:Mapping|CO\s*PO|$)', co_text, re.DOTALL | re.IGNORECASE)
                if statement_match:
                    statement_text = statement_match.group(1)
                    # Extract lines that look like outcomes (start with capital, long enough)
                    co_num = 1
                    current_outcome = ""
                    
                    for line in statement_text.split('\n'):
                        line = line.strip()
                        # Skip empty lines and noise
                        if not line or len(line) < 10:
                            continue
                        # Skip lines that look like table headers or mapping data
                        if re.match(r'^(CO\s*No|PO|PSO|BTL|\d+\s*$)', line, re.IGNORECASE):
                            continue
                        # Skip lines that are mostly numbers (mapping rows)
                        if re.match(r'^[\d\s,.-]+$', line):
                            continue
                        
                        # Check if this starts with a capital letter (new outcome)
                        if re.match(r'^[A-Z]', line):
                            # Save previous outcome if exists
                            if current_outcome:
                                bloom_level = self.text_processor.classify_bloom_level(current_outcome)
                                outcomes.append({
                                    'code': f'CO{co_num}',
                                    'description': current_outcome.strip(),
                                    'bloom_level': bloom_level
                                })
                                co_num += 1
                            current_outcome = line
                        else:
                            # Continuation of previous outcome
                            current_outcome += " " + line
                    
                    # Don't forget the last outcome
                    if current_outcome and len(current_outcome) >= 20:
                        bloom_level = self.text_processor.classify_bloom_level(current_outcome)
                        outcomes.append({
                            'code': f'CO{co_num}',
                            'description': current_outcome.strip(),
                            'bloom_level': bloom_level
                        })
        
        # Fallback: Pattern for CO1:, CO2:, etc.
        if not outcomes:
            pattern = r'(CO\d+)[:\s]+([A-Za-z][^\n]+?)(?=CO\d+|Mapping|$)'
            matches = re.finditer(pattern, text, re.IGNORECASE)
            
            for match in matches:
                co_code = match.group(1).upper()
                co_text = match.group(2).strip()
                
                # Skip if it looks like mapping numbers
                if re.match(r'^[\d\s,.-]+$', co_text):
                    continue
                
                bloom_level = self.text_processor.classify_bloom_level(co_text)
                
                outcomes.append({
                    'code': co_code,
                    'description': co_text,
                    'bloom_level': bloom_level
                })
            
        return outcomes
        
    def _extract_units(self, text: str) -> List[Dict[str, Any]]:
        """Extract unit-wise syllabus"""
        units = []
        
        # Multiple patterns for different unit formats, including OCR-tolerant versions
        patterns = [
            # Format: "Unit No I Title Hours" or "Unit No 1 Title Hours" (same line)
            r'Unit\s*No\.?\s*([IVX]+|\d+)\s+(.+?)\s+(\d+)\s*H[o0]u?rs?',
            # OCR-tolerant: "UNIT NO.04" or "UNIT NOS" (S misread from 5) followed by title and hours
            r'UNIT\s*NO\.?\s*([0-9]+|[O0][1-9]|[1-9][0-9]?)\s+(.+?)\s+(\d+)\s*H[o0]w?u?rs?',
            # Format: "Unit No 1 Title" followed by "Hours X" on next line (flattened table)
            r'Unit\s*No\.?\s*(\d+)\s+([A-Za-z][A-Za-z\s&,-]+?)(?=\nHours)',
            # Format: "Unit 1: Title" or "Unit I: Title"
            r'Unit\s+(\d+|[IVX]+)[:\s]+(.+?)(?=\n)',
        ]
        
        # Try first pattern (Unit No format with hours on same line)
        pattern = patterns[0]
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        
        if matches:
            for i, match in enumerate(matches):
                unit_num = match.group(1)
                title = match.group(2).strip()
                hours = int(match.group(3)) if match.group(3) else 0
                
                # Find content between this unit and next unit (or end)
                start = match.end()
                if i + 1 < len(matches):
                    end = matches[i + 1].start()
                else:
                    # Find next section marker
                    next_section = re.search(r'\n(Textbooks?|References?|Assessment)', text[start:], re.IGNORECASE)
                    end = start + next_section.start() if next_section else len(text)
                
                unit_content = text[start:end].strip()
                
                # Extract topics - lines that contain actual content
                topics = []
                for line in unit_content.split('\n'):
                    line = line.strip()
                    # Skip empty lines, page headers, and CO/BTL references
                    if line and len(line) > 10 and not re.match(r'^[\d\s,]+$', line):
                        if not re.match(r'^(Unit|CO|BTL|M\.Sc|Vishwakarma)', line, re.IGNORECASE):
                            topics.append(line)
                
                units.append({
                    'unit_number': unit_num,
                    'title': title,
                    'topics': topics[:10],  # Limit topics per unit
                    'hours': hours
                })
        else:
            # Try second pattern (flattened table: "Unit No X Title" then "Hours Y" on next line)
            pattern = patterns[1]
            header_matches = list(re.finditer(pattern, text, re.IGNORECASE))
            
            if header_matches:
                for i, match in enumerate(header_matches):
                    unit_num = match.group(1)
                    title = match.group(2).strip()
                    
                    # Look for Hours on next line
                    hours_search_start = match.end()
                    hours_match = re.search(r'Hours\s*(\d+)', text[hours_search_start:hours_search_start+50], re.IGNORECASE)
                    hours = int(hours_match.group(1)) if hours_match else 0
                    
                    # Find content between this unit's Hours and next unit (or end)
                    if hours_match:
                        start = hours_search_start + hours_match.end()
                    else:
                        start = match.end()
                    
                    if i + 1 < len(header_matches):
                        end = header_matches[i + 1].start()
                    else:
                        # Find next section marker
                        next_section = re.search(r'\n(Textbooks?|References?|Assessment|CO\n)', text[start:], re.IGNORECASE)
                        end = start + next_section.start() if next_section else len(text)
                    
                    unit_content = text[start:end].strip()
                    
                    # Extract topics - lines that contain actual content
                    topics = []
                    for line in unit_content.split('\n'):
                        line = line.strip()
                        # Skip empty lines, page headers, CO/BTL references, and noise
                        if line and len(line) > 10 and not re.match(r'^[\d\s,]+$', line):
                            if not re.match(r'^(Unit|CO|BTL|M\.Sc|Vishwakarma|Hours|BTECH|Pattern)', line, re.IGNORECASE):
                                topics.append(line)
                    
                    units.append({
                        'unit_number': unit_num,
                        'title': title,
                        'topics': topics[:10],
                        'hours': hours
                    })
            else:
                # Try third pattern (Unit 1: format)
                pattern = patterns[2]
                matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
                
                for match in matches:
                    unit_num = match.group(1)
                    unit_content = match.group(2).strip()
                    
                    lines = unit_content.split('\n')
                    title = lines[0].strip() if lines else ""
                    
                    hours_match = re.search(r'(\d+)\s*(?:hours?|hrs?)', unit_content, re.IGNORECASE)
                    hours = int(hours_match.group(1)) if hours_match else 0
                    
                    topics = [line.strip() for line in lines[1:] if line.strip() and len(line.strip()) > 10]
                    
                    units.append({
                        'unit_number': unit_num,
                        'title': title,
                        'topics': topics[:10],
                        'hours': hours
                    })
            
        return units
        
    def _extract_assessment(self, text: str) -> Dict[str, Any]:
        """Extract assessment pattern"""
        assessment = {}
        
        # Common assessment components
        components = ['internal', 'external', 'midterm', 'final', 'assignment', 'quiz', 'project', 'lab']
        
        for component in components:
            pattern = rf'{component}[:\s]+(\d+)%?'
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                assessment[component] = int(match.group(1))
                
        return assessment
        
    def _extract_references(self, text: str) -> List[str]:
        """Extract reference books and materials"""
        import re
        
        # Find the references section
        section_pattern = r'(?:Reference\s*Books?|References?|Text\s*Books?|Bibliography)[:\s]*(.+?)(?=\n\n[A-Z]|\nCourse\s*Outcome|\nMapping|\nUnit\s*No|\nAssessment|$)'
        match = re.search(section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if not match:
            return []
        
        ref_text = match.group(1).strip()
        references = []
        
        # Check if text has inline numbered items (1. ... 2. ... 3. ...)
        inline_pattern = r'(\d+)\.\s+'
        inline_matches = list(re.finditer(inline_pattern, ref_text))
        
        # Count how many inline number patterns we have
        if len(inline_matches) >= 2:
            # Multiple numbered items found - extract between them
            for idx, match_item in enumerate(inline_matches):
                start = match_item.end()
                if idx + 1 < len(inline_matches):
                    end = inline_matches[idx + 1].start()
                else:
                    end = len(ref_text)
                
                ref_clean = ref_text[start:end].strip()
                ref_clean = re.sub(r'\s+', ' ', ref_clean)
                if ref_clean and len(ref_clean) > 10:
                    references.append(ref_clean)
        
        # If inline method found items, we're done
        if references:
            return references
        
        # Method 2: Try newline-separated numbered items (1. Book\n2. Book)
        numbered_pattern = r'(?:^|\n)\s*(\d+)[.\)]\s*(.+?)(?=\n\s*\d+[.\)]|$)'
        numbered_matches = re.findall(numbered_pattern, ref_text, re.DOTALL)
        
        if len(numbered_matches) >= 2:
            for num, ref in numbered_matches:
                ref_clean = ref.strip().replace('\n', ' ')
                ref_clean = re.sub(r'\s+', ' ', ref_clean)
                if ref_clean and len(ref_clean) > 10:
                    references.append(ref_clean)
        
        if references:
            return references
        
        # Method 3: Try bulleted items
        bullet_pattern = r'[•\*]\s*(.+?)(?=[•\*]|\n|$)'
        bullet_matches = re.findall(bullet_pattern, ref_text, re.DOTALL)
        for ref in bullet_matches:
            ref_clean = ref.strip().replace('\n', ' ')
            ref_clean = re.sub(r'\s+', ' ', ref_clean)
            if ref_clean and len(ref_clean) > 10:
                references.append(ref_clean)
        
        if references:
            return references
        
        # Method 4: Single reference - just return the whole text if it looks like a reference
        if len(ref_text) > 20 and not re.match(r'^[\d\s,.-]+$', ref_text):
            return [re.sub(r'\s+', ' ', ref_text.strip())]
        
        return references
        
    def _extract_co_po_mapping(self, text: str) -> Dict[str, Dict[str, int]]:
        """Extract CO-PO mapping matrix"""
        mapping = {}
        
        # Look for mapping table
        # This is a simplified version - real implementation would need table parsing
        pattern = r'(CO\d+).*?(PO\d+)[:\s]+(\d+)'
        matches = re.finditer(pattern, text, re.IGNORECASE)
        
        for match in matches:
            co = match.group(1).upper()
            po = match.group(2).upper()
            level = int(match.group(3))
            
            if co not in mapping:
                mapping[co] = {}
            mapping[co][po] = level
            
        return mapping
