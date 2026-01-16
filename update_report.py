
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # helper for style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Project Report\n\n")
    run.font.size = Pt(24)
    run.font.bold = True
    
    run2 = title_para.add_run("Syllabus and Curriculum Design Optimizer (SCDO)\n\n")
    run2.font.size = Pt(18)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("\n" * 5)
    
    details_para = doc.add_paragraph()
    details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details_run = details_para.add_run("Updated: January 2026\n")
    details_run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "The Syllabus and Curriculum Design Optimizer (SCDO) is a cutting-edge web application "
        "designed to revolutionize the way academic syllabi are created, analyzed, and optimized. "
        "Leveraging the power of IBM Granite AI models and State-of-the-Art (SOTA) Natural Language "
        "Processing, SCDO assists faculty and curriculum designers in crafting high-quality, industry-relevant courses."
    )
    
    # 2. Key Features
    doc.add_heading('2. Key Features', level=1)
    
    doc.add_heading('2.1 AI-Powered Syllabus Generation', level=2)
    doc.add_paragraph(
        "Users can generate complete, structured syllabi from minimal inputs (Title, Credits, Course Code). "
        "The system utilizes IBM Granite LLMs to produce detailed unit content, learning outcomes (mapped to Bloom's Taxonomy), "
        "and assessment patterns."
    )

    doc.add_heading('2.2 Intelligent Gap Analysis', level=2)
    doc.add_paragraph(
        "The Analyze module parses existing syllabus files (PDF, DOCX, TXT) to identify weaknesses. Features include:"
    )
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Automated Quality Scoring (0-100) based on completeness and structure.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Bloom's Taxonomy Distribution visualization.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Content Depth Analysis showing distribution of Basic vs. Advanced units.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Redundancy Detection identifying overlapping topics.")

    doc.add_heading('2.3 Optimization & Enhancement', level=2)
    doc.add_paragraph(
        "The Optimize module provides actionable recommendations to improve course quality:"
    )
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Rebalancing suggestions for Bloom's levels.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Modern topic suggestions to ensure industry relevance.")
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.add_run("Course Outcome to Program Outcome (CO-PO) mapping.")
    
    doc.add_heading('2.4 Professional Reports', level=2)
    doc.add_paragraph(
        "Updated export functionality now allows generating comprehensive PDF reports that include "
        "not just the syllabus content but also the detailed Analysis Report (charts, scores, and gaps)."
    )
    
    # 3. Technical Architecture
    doc.add_heading('3. Technical Architecture', level=1)
    
    doc.add_heading('3.1 Frontend', level=2)
    doc.add_paragraph(
        "Built with React.js and Vite, providing a fast, responsive Single Page Application (SPA). "
        "Recent updates include a comprehensive Dark Mode (using CSS variables) and micro-animations for a "
        "premium user experience."
    )
    
    doc.add_heading('3.2 Backend', level=2)
    doc.add_paragraph(
        "Powered by FastAPI (Python), handling complex parsing, AI orchestration, and file generation. "
        "Key components include the PDFExporter (ReportLab) and integration with IBM WatsonX AI services."
    )
    
    # 4. Recent Improvements
    doc.add_heading('4. Recent Improvements', level=1)
    doc.add_paragraph("Several critical enhancements have been made in the latest sprint:")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Enhancement'
    hdr_cells[1].text = 'Description'
    
    updates = [
        ("Dark Mode Support", "Full theme compatibility across all pages (Analyze, Generate, etc.) using CSS variables."),
        ("PDF Analysis Export", "Updated report generator to include visual charts and quality scores in PDF output."),
        ("UI Layout Fixes", "Resolved overflow issues in Content Depth and Modern Topic cards."),
        ("Robust Backend", "Improved error handling and self-contained export logic (removing dependencies on system tools).")
    ]
    
    for item, desc in updates:
        row_cells = table.add_row().cells
        row_cells[0].text = item
        row_cells[1].text = desc

    # 5. Conclusion
    doc.add_heading('5. Conclusion', level=1)
    doc.add_paragraph(
        "The SCDO platform has reached a mature stage with robust analysis and generation capabilities. "
        "The integration of modern UI practices and comprehensive reporting features ensures it meets the "
        "needs of academic institutions aiming for curriculum excellence."
    )

    doc.save("d:\\Syllabus Optimizer\\Project_Report_Updated.docx")
    print("Report generated: d:\\Syllabus Optimizer\\Project_Report_Updated.docx")

if __name__ == "__main__":
    create_report()
