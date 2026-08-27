from fastapi import APIRouter, HTTPException, Depends
from fastapi.responses import FileResponse
import tempfile
import logging

from app.schemas import OptimizeRequest
from app.dependencies import verify_api_key, get_components
from src.utils.logging_utils import setup_logger
from src.export.excel_exporter import ExcelExporter

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    from src.export.latex_exporter import LaTeXExporter
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False

logger = setup_logger("scdo_api", log_file="logs/api.log")

router = APIRouter(dependencies=[Depends(verify_api_key)])

@router.post("/api/export/pdf")
async def export_pdf(request: OptimizeRequest, comps=Depends(get_components)):
    """Export syllabus and analysis to PDF"""
    try:
        syllabus_data = request.syllabus_data
        analysis_data = request.analysis_data
        
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        success = comps.pdf_exporter.export(
            syllabus_data, 
            pdf_path, 
            include_mapping=True,
            analysis_data=analysis_data
        )
        
        if not success:
            raise HTTPException(status_code=500, detail="PDF export failed")
        
        filename = f"{syllabus_data.get('course_code', 'syllabus')}_analysis.pdf" if analysis_data else f"{syllabus_data.get('course_code', 'syllabus')}.pdf"
        
        return FileResponse(pdf_path, media_type="application/pdf", filename=filename)
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        raise HTTPException(status_code=500, detail="PDF export failed.")


@router.post("/api/export/latex-pdf")
async def export_latex_pdf(request: OptimizeRequest):
    """Export syllabus to PDF using LaTeX"""
    if not LATEX_AVAILABLE:
        raise HTTPException(status_code=503, detail="LaTeX exporter not available.")
    try:
        syllabus_data = request.syllabus_data
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        latex_exporter = LaTeXExporter()
        result_path = latex_exporter.export_pdf(syllabus_data, pdf_path)
        
        if result_path.endswith('.tex'):
            media_type = "text/x-tex"
            filename = f"{syllabus_data.get('course_code', 'syllabus')}.tex"
        else:
            media_type = "application/pdf"
            filename = f"{syllabus_data.get('course_code', 'syllabus')}_latex.pdf"
        
        return FileResponse(result_path, media_type=media_type, filename=filename)
    except Exception as e:
        logger.error(f"LaTeX PDF export failed: {e}")
        raise HTTPException(status_code=500, detail="LaTeX PDF export failed.")


@router.post("/api/export/excel")
async def export_excel(request: OptimizeRequest):
    """Export syllabus to Excel"""
    try:
        syllabus_data = request.syllabus_data
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            output_path = tmp.name
            
        excel_exporter = ExcelExporter()
        success = excel_exporter.export_complete_syllabus(
            syllabus_data, output_path, include_mapping=True, include_rubrics=True
        )
        
        if not success:
            raise HTTPException(status_code=500, detail="Excel export failed")
            
        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            filename=f"{syllabus_data.get('course_code', 'syllabus')}.xlsx"
        )
    except Exception as e:
        logger.error(f"Excel export failed: {e}")
        raise HTTPException(status_code=500, detail="Excel export failed.")


@router.post("/api/export/word")
async def export_word(request: OptimizeRequest):
    """Export syllabus to Word document"""
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=501, detail="Word export not available (python-docx missing).")
    try:
        syllabus_data = request.syllabus_data
        analysis_data = request.analysis_data
        
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        if analysis_data:
            if analysis_data.get('co_po_mapping') and not syllabus_data.get('co_po_mapping'):
                syllabus_data['co_po_mapping'] = analysis_data['co_po_mapping']
            if analysis_data.get('bloom_analysis') and not syllabus_data.get('bloom_analysis'):
                syllabus_data['bloom_analysis'] = analysis_data['bloom_analysis']
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            output_path = tmp.name
        
        doc = Document()
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if syllabus_data.get('university_name'):
            uni = doc.add_paragraph(syllabus_data.get('university_name'))
            uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = uni.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            
        if syllabus_data.get('faculty_name'):
            fac = doc.add_paragraph(syllabus_data.get('faculty_name'))
            fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fac.runs[0]
            run.font.size = Pt(12)
            
        if syllabus_data.get('program') or syllabus_data.get('department'):
            prog_text = syllabus_data.get('program') or syllabus_data.get('department', '')
            prog = doc.add_paragraph(prog_text)
            prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = prog.runs[0]
            run.font.size = Pt(11)
            run.font.italic = True
            
        doc.add_paragraph()
        
        course_code = syllabus_data.get('course_code', '')
        course_title = syllabus_data.get('course_title', 'Course Syllabus')
        title_text = f"{course_code}: {course_title}" if course_code else course_title
        
        title_para = doc.add_paragraph(title_text)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph()
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        row0 = table.rows[0]
        row0.cells[0].text = f"Course Type: {syllabus_data.get('course_type', 'DSC')}"
        row0.cells[1].text = f"Semester: {syllabus_data.get('semester', 'I')}"
        
        credits = syllabus_data.get('credits', '3-0-0')
        year = syllabus_data.get('year', '')
        row1 = table.rows[1]
        row1.cells[0].text = f"Credits (L-T-P): {credits}"
        row1.cells[1].text = f"Academic Year: {year}" if year else ""
        
        for row in table.rows:
            for cell in row.cells:
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        if syllabus_data.get('overview'):
            h = doc.add_heading('Course Overview', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            doc.add_paragraph(syllabus_data['overview'])
            doc.add_paragraph()
            
        if syllabus_data.get('learning_outcomes'):
            h = doc.add_heading('Course Learning Outcomes (COs)', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Grid Table 4 Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'CO'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = "Bloom's Level"
            
            for outcome in syllabus_data['learning_outcomes']:
                row_cells = table.add_row().cells
                row_cells[0].text = outcome.get('code', '')
                row_cells[1].text = outcome.get('description', '')
                row_cells[2].text = outcome.get('bloom_level', '').capitalize()
            
            table.autofit = False
            table.columns[0].width = Inches(0.8)
            table.columns[1].width = Inches(3.8)
            table.columns[2].width = Inches(1.4)
            doc.add_paragraph()
            
        if syllabus_data.get('units'):
            h = doc.add_heading('Course Content', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            for unit in syllabus_data['units']:
                unit_num = unit.get('unit_number', '')
                title = unit.get('title', 'Untitled')
                hours = unit.get('hours', 0)
                
                utable = doc.add_table(rows=1, cols=2)
                utable.allow_autofit = False
                utable.columns[0].width = Inches(5.0)
                utable.columns[1].width = Inches(1.0)
                
                c1 = utable.rows[0].cells[0]
                p = c1.paragraphs[0]
                run = p.add_run(f"Unit {unit_num}: {title}")
                run.bold = True
                run.font.size = Pt(11)
                
                c2 = utable.rows[0].cells[1]
                p = c2.paragraphs[0]
                run = p.add_run(f"{hours} Hours")
                run.bold = True
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                topics = unit.get('topics', [])
                if topics:
                    topics_text = ", ".join(topics)
                    p = doc.add_paragraph(topics_text)
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(12)
                else:
                    doc.add_paragraph()

        if syllabus_data.get('references'):
             h = doc.add_heading('References', level=1)
             h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
             refs = syllabus_data.get('references', {})
             
             if isinstance(refs, dict):
                 if refs.get('textbooks'):
                     doc.add_heading('Textbooks:', level=2)
                     for t in refs['textbooks']: doc.add_paragraph(t, style='List Bullet')
                 if refs.get('references'):
                     doc.add_heading('Reference Books:', level=2)
                     for r in refs['references']: doc.add_paragraph(r, style='List Bullet')
                 if refs.get('online_resources'):
                     doc.add_heading('Online Resources:', level=2)
                     for r in refs['online_resources']: doc.add_paragraph(r, style='List Bullet')
             elif isinstance(refs, list):
                 for r in refs: doc.add_paragraph(r, style='List Bullet')

        doc.save(output_path)
        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{syllabus_data.get('course_code', 'syllabus')}.docx"
        )
    except Exception as e:
        logger.error(f"Word export failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail="Word export failed.")
