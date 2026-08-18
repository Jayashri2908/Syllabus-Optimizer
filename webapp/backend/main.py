"""
FastAPI Backend for SCDO
Main API server for syllabus optimization
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks, Request, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path
import tempfile
import os
import re
import secrets

# Maximum upload file size in bytes (default: 50 MB)
MAX_UPLOAD_SIZE = int(os.getenv("MAX_UPLOAD_SIZE", 50 * 1024 * 1024))

# API key for endpoint protection (set API_KEY env var to enable, leave empty to disable)
API_KEY = os.getenv("API_KEY", "")

# Force HF to use local files only and skip version checks
os.environ['HF_HUB_OFFLINE'] = '1'
os.environ['TRANSFORMERS_OFFLINE'] = '1'

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv(Path(__file__).parent.parent.parent / ".env")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Word export will not be available.")

# Import SCDO modules
import sys
sys.path.append(str(Path(__file__).parent.parent.parent))

from src.analysis.syllabus_parser import SyllabusParser
from src.analysis.gap_analyzer import GapAnalyzer
from src.analysis.outcome_extractor import OutcomeExtractor
from src.optimization.bloom_mapper import BloomMapper
from src.optimization.content_optimizer import ContentOptimizer
from src.optimization.objectives_optimizer import ObjectivesOptimizer
from src.optimization.reference_suggester import ReferenceSuggester
from src.generation.syllabus_generator import SyllabusGenerator
from src.utils.mock_services import MockContentOptimizer, MockBloomMapper, MockGapAnalyzer # Import Mocks
from src.analysis.rag_analyzer import RAGAwareAnalyzer # RAG
from src.mapping.co_po_mapper import COPOMapper
from src.export.pdf_exporter import PDFExporter
from src.export.excel_exporter import ExcelExporter
try:
    from src.export.latex_exporter import LaTeXExporter
    LATEX_AVAILABLE = True
except ImportError:
    LATEX_AVAILABLE = False
    logging.warning("LaTeX exporter not available")
from src.ibm.local_storage import LocalStorage
from src.utils.logging_utils import setup_logger

# Setup logging
logger = setup_logger("scdo_api", log_file="logs/api.log")

# Initialize FastAPI app
app = FastAPI(
    title="Syllabus and Curriculum Design Optimizer API",
    description="AI-powered syllabus analysis, optimization, and generation",
    version="1.0.0"
)

# Cors middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:3000").split(","),  # Restrict to known origins; set CORS_ORIGINS env var for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

logger.info("="*50)
logger.info("  SCDO BACKEND SERVER - AI-POWERED ONLY  ")
logger.info("  IBM Granite Integration Required  ")
logger.info("="*50)

# Optional API key authentication
# Set API_KEY env var to enable; leave empty to allow all requests (dev mode)
async def verify_api_key(request: Request):
    """Verify API key from X-API-Key header or api_key query param.
    Skips auth if API_KEY env var is not set (development mode)."""
    if not API_KEY:
        return  # Auth disabled in dev mode
    
    key = request.headers.get("X-API-Key") or request.query_params.get("api_key")
    if key != API_KEY:
        raise HTTPException(
            status_code=401,
            detail="Invalid or missing API key. Set X-API-Key header or api_key query param."
        )

# Initialize components
parser = None
gap_analyzer = None
outcome_extractor = None
bloom_mapper = None
content_optimizer = None
syllabus_generator = None
co_po_mapper = None
pdf_exporter = None
local_storage = None
objectives_optimizer = None # NEW
reference_suggester = None # NEW

# Initialize critical components (Local)
try:
    parser = SyllabusParser()
    outcome_extractor = OutcomeExtractor()
    co_po_mapper = COPOMapper()
    pdf_exporter = PDFExporter()
    local_storage = LocalStorage()  # FREE local storage instead of cloud
    logger.info("Critical components initialized successfully")
except Exception as e:
    logger.error(f"Critical component initialization failed: {e}")
    # We continue, but API will likely fail on specific endpoints

# Initialize AI components (IBM Cloud) - AI-ONLY MODE
bloom_mapper_initialized = False
content_optimizer_initialized = False
try:
    syllabus_generator = SyllabusGenerator()
    gap_analyzer = GapAnalyzer()
    bloom_mapper = BloomMapper()  # Initialize real BloomMapper
    content_optimizer = ContentOptimizer()  # Initialize real ContentOptimizer
    objectives_optimizer = ObjectivesOptimizer()  # Initialize
    reference_suggester = ReferenceSuggester()  # Initialize
    bloom_mapper_initialized = True
    content_optimizer_initialized = True
    logger.info("[OK] AI components initialized successfully (AI models active)")
except Exception as e:
    logger.error(f"[ERROR] AI component initialization failed: {e}")
    logger.error("[WARNING] IBM Granite credentials required! Run: python setup_credentials.py")
    logger.warning("Generation endpoint will return error until credentials are configured.")
    
    # Try to initialize RAG analyzer as fallback for analysis
    try:
        gap_analyzer = RAGAwareAnalyzer()
        logger.info("Initialized RAGAwareAnalyzer for analysis")
    except Exception as rag_err:
        logger.warning(f"RAG init failed: {rag_err}, using basic mock")
        gap_analyzer = MockGapAnalyzer()

# Ensure bloom_mapper and content_optimizer are always initialized
if not bloom_mapper_initialized:
    bloom_mapper = MockBloomMapper()
if not content_optimizer_initialized:
    content_optimizer = MockContentOptimizer()



# Pydantic models
class GenerateRequest(BaseModel):
    # Institution Details
    university_name: str = ""
    faculty_name: str = ""
    department: str = ""
    
    # Course Details
    course_title: str
    course_code: str
    course_type: str = "DSC"  # DSC, DSE, GEC, SEC, etc.
    credits: str
    semester: str = "I"
    program: str = ""
    year: str = ""
    course_level: str = "intermediate"
    
    # Content
    program_outcomes: List[str]
    keywords: List[str] = []
    unit_topics: List[Dict[str, Any]] = []
    
    # References
    textbooks: List[str] = []
    references: List[str] = []
    online_resources: List[str] = []
    
    # Settings
    domain: str = "engineering"
    num_units: int = 5
    num_outcomes: int = 5


class OptimizeRequest(BaseModel):
    syllabus_data: Dict[str, Any]
    optimization_goals: List[str] = []
    analysis_data: Optional[Dict[str, Any]] = None  # Added field


class MapRequest(BaseModel):
    course_outcomes: List[Dict[str, str]]
    program_outcomes: Optional[List[str]] = None
    domain: str = "engineering"


# API Endpoints
@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "Syllabus and Curriculum Design Optimizer API",
        "version": "1.0.0",
        "status": "operational"
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "SCDO API"}


@app.post("/api/upload")
async def upload_syllabus(file: UploadFile = File(...), auth=Depends(verify_api_key)):
    """
    Upload and parse syllabus file
    
    Accepts: PDF, DOCX, TXT files
    Returns: Parsed syllabus structure
    """
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.txt']
        file_ext = Path(file.filename).suffix.lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Sanitize filename to prevent path traversal
        safe_filename = re.sub(r'[^\w\s\-.]', '_', Path(file.filename).stem) + file_ext
        
        # Read and validate file size
        content = await file.read()
        if len(content) > MAX_UPLOAD_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {MAX_UPLOAD_SIZE // (1024*1024)} MB"
            )
        
        # Save uploaded file temporarily
        if not local_storage:
             # Fallback to temp file if local_storage failed init (though it shouldn't)
             with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
        else:
             tmp_path = await local_storage.save_upload(file, content=content, filename=safe_filename)
        
        try:
            # Parse syllabus
            if not parser:
                raise HTTPException(status_code=503, detail="Syllabus Parser unavailable")
                
            syllabus_data = parser.parse_file(tmp_path)
            
            logger.info(f"Successfully parsed syllabus: {file.filename} ({len(syllabus_data.get('units', []))} units, {len(syllabus_data.get('learning_outcomes', []))} outcomes)")
            
            return {
                "success": True,
                "filename": file.filename,
                "data": syllabus_data
            }
            
        finally:
            # Clean up temp file if it was a temp file (local_storage files persist)
            if not local_storage and os.path.exists(tmp_path):
                os.unlink(tmp_path)
            
    except Exception as e:
        logger.error(f"Upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/analyze")
async def analyze_syllabus(syllabus_data: Dict[str, Any], auth=Depends(verify_api_key)):
    """
    Analyze syllabus for gaps and issues
    
    Returns: Comprehensive gap analysis report
    """
    if not gap_analyzer:
        raise HTTPException(status_code=503, detail="Gap Analyzer service unavailable (check IBM credentials)")
        
    try:
        logger.info(f"Analyzing syllabus for course: {syllabus_data.get('course_title', 'Unknown')}")
        analysis = gap_analyzer.analyze(syllabus_data)
        logger.info(f"Analysis complete for {syllabus_data.get('course_title', 'Unknown')}")
        
        return {
            "success": True,
            "analysis": analysis
        }
        
    except Exception as e:
        logger.error(f"Analysis failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/optimize")
async def optimize_syllabus(request: OptimizeRequest, auth=Depends(verify_api_key)):
    """
    Perform a complete optimization of the syllabus using a unified pipeline.
    Returns: Both original and optimized syllabus for side-by-side comparison.
    """
    try:
        original_syllabus = request.syllabus_data
        course_title = original_syllabus.get('course_title', 'Unknown')
        logger.info(f"Starting unified optimization for course: {course_title}")
        
        # Determine if we should use fallback AI analysis (incomplete parsing)
        outcomes = original_syllabus.get('learning_outcomes', [])
        units = original_syllabus.get('units', [])
        raw_text = original_syllabus.get('raw_text', '')
        
        has_enough_data = (outcomes and len(outcomes) > 0) or (units and len(units) > 0)
        
        if not has_enough_data and raw_text:
            logger.info(f"Incomplete parsing for '{course_title}' - using AI to optimize from raw content")
            # In this case, we treat original_syllabus as the starting point but the prompt handles the rest
            pass

        # Call unified optimization pipeline
        if not content_optimizer:
            raise HTTPException(status_code=503, detail="Content Optimizer service unavailable")
            
        logger.info(f"Calling LLM for optimization of '{course_title}'...")
        optimization_result = content_optimizer.optimize_full_syllabus(original_syllabus)
        logger.info(f"LLM optimization complete for '{course_title}'")
        
        # New optimized structure
        optimized_syllabus = optimization_result.get('optimized_syllabus', original_syllabus)
        
        # Post-processing: Apply compliance checks on the OPTIMIZED syllabus
        logger.info(f"Running post-processing compliance checks for '{course_title}'")
        nep_2020_compliance = None
        try:
            from src.validation.nep_2020_validator import NEP2020Validator
            nep_validator = NEP2020Validator()
            nep_2020_compliance = nep_validator.validate(optimized_syllabus)
            logger.info(f"NEP 2020 validation complete for '{course_title}'")
        except Exception as e:
            logger.error(f"NEP 2020 validation failed: {e}")
            
        accreditation_compliance = None
        try:
            from src.validation.accreditation_checker import AccreditationChecker
            accred_checker = AccreditationChecker()
            accreditation_compliance = {
                'nba': accred_checker.check_nba_compliance(optimized_syllabus),
                'naac': accred_checker.check_naac_compliance(optimized_syllabus)
            }
            logger.info(f"Accreditation checks complete for '{course_title}'")
        except Exception as e:
            logger.error(f"Accreditation check failed: {e}")

        # Post-processing: CO-PO Mapping on the OPTIMIZED syllabus
        co_po_mapping = None
        if co_po_mapper:
            try:
                logger.info(f"Performing CO-PO mapping for '{course_title}'")
                co_po_mapping = co_po_mapper.map_co_to_po(
                    course_outcomes=optimized_syllabus.get('learning_outcomes', []),
                    program_outcomes=original_syllabus.get('program_outcomes', [
                        'PO1', 'PO2', 'PO3', 'PO4', 'PO5', 'PO6', 'PO7', 'PO8', 'PO9', 'PO10', 'PO11', 'PO12'
                    ]),
                    domain=optimized_syllabus.get('domain', 'engineering')
                )
                logger.info(f"CO-PO mapping complete for '{course_title}'")
            except Exception as e:
                logger.error(f"CO-PO mapping on optimized syllabus failed: {e}")

        logger.info(f"Optimization pipeline successfully finished for '{course_title}'")
        return {
            'success': True,
            'original_syllabus': original_syllabus,
            'optimized_syllabus': optimized_syllabus,
            'optimization': {
                'changes_summary': optimization_result.get('changes_summary', []),
                'bloom_distribution': optimization_result.get('bloom_distribution', {}),
                'rationale': optimization_result.get('rationale', ''),
                'industry_relevance_score': optimization_result.get('industry_relevance_score', 0),
                'prerequisite_rationale': optimization_result.get('prerequisite_rationale', ''),
                'nep_2020_compliance': nep_2020_compliance,
                'accreditation_compliance': accreditation_compliance,
                'co_po_mapping': co_po_mapping
            }
        }
        
    except Exception as e:
        logger.error(f"Optimization pipeline failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Optimization failed: {str(e)}")



@app.post("/api/generate")
async def generate_syllabus(request: GenerateRequest, auth=Depends(verify_api_key)):
    """
    Generate new syllabus from minimal inputs using AI
    
    Requires: IBM Granite credentials configured
    Returns: Complete generated syllabus with quality scoring
    """
    if not syllabus_generator:
        raise HTTPException(
            status_code=503,
            detail="AI Syllabus Generator unavailable. IBM Granite credentials required. Run: python setup_credentials.py"
        )
    
    try:
        syllabus = syllabus_generator.generate(
            course_title=request.course_title,
            course_code=request.course_code,
            credits=request.credits,
            program=request.program,
            year=request.year,
            course_level=request.course_level,
            program_outcomes=request.program_outcomes,
            keywords=request.keywords,
            unit_topics=request.unit_topics,
            textbooks=request.textbooks,
            references=request.references,
            online_resources=request.online_resources,
            domain=request.domain,
            num_units=request.num_units,
            num_outcomes=request.num_outcomes,
            use_chained_generation=False  # Use upgraded SyllabusGenerator with structured JSON
        )
        
        # Add institution details to syllabus
        syllabus['university_name'] = request.university_name
        syllabus['faculty_name'] = request.faculty_name
        syllabus['department'] = request.department
        syllabus['course_type'] = request.course_type
        syllabus['semester'] = request.semester
        
        # DEBUG: Log unit/topic structure
        if syllabus.get('units'):
            first_unit = syllabus['units'][0]
            logger.info(f"DEBUG: Unit type: {type(first_unit)}")
            if isinstance(first_unit, dict) and 'topics' in first_unit:
                logger.info(f"DEBUG: Topics type: {type(first_unit['topics'])}")
                if first_unit['topics']:
                    logger.info(f"DEBUG: First topic type: {type(first_unit['topics'][0])}")
                    logger.info(f"DEBUG: First topic value: {first_unit['topics'][0]}")
        
        return {
            "success": True,
            "syllabus": syllabus
        }
        
    except Exception as e:
        logger.error(f"Generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/map-outcomes")
async def map_outcomes(request: MapRequest, auth=Depends(verify_api_key)):
    """
    Perform CO-PO mapping
    
    Returns: CO-PO mapping matrix and validation
    """
    try:
        # Generate mapping
        mapping = co_po_mapper.map_co_to_po(
            course_outcomes=request.course_outcomes,
            program_outcomes=request.program_outcomes,
            domain=request.domain
        )
        
        # Generate formatted matrix
        matrix_text = co_po_mapper.generate_mapping_matrix(mapping)
        
        # Validate mapping
        validation = co_po_mapper.validate_mapping(mapping)
        
        return {
            "success": True,
            "mapping": mapping,
            "matrix": matrix_text,
            "validation": validation
        }
        
    except Exception as e:
        logger.error(f"Mapping failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export/pdf")
async def export_pdf(request: OptimizeRequest, auth=Depends(verify_api_key)):
    """
    Export syllabus and analysis to PDF
    
    Returns: PDF file
    """
    try:
        # Extract syllabus and analysis data
        syllabus_data = request.syllabus_data
        analysis_data = request.analysis_data
        
        # Handle nested structure - if data is wrapped, unwrap it
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            # Frontend sent upload response structure
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            # Frontend sent optimize response structure
            syllabus_data = syllabus_data['syllabus']
        
        # Log what we received for debugging
        logger.info(f"Exporting PDF for course: {syllabus_data.get('course_title', 'N/A')}")
        if analysis_data:
            logger.info(f"Analysis data keys: {list(analysis_data.keys())}")
            if analysis_data.get('ai_analysis'):
                logger.info(f"AI analysis found: {len(analysis_data['ai_analysis'])} chars")
            if analysis_data.get('sequence_optimization'):
                logger.info(f"Sequence optimization found: {analysis_data['sequence_optimization'].keys() if isinstance(analysis_data['sequence_optimization'], dict) else 'string'}")
        else:
            logger.info("No analysis_data provided")
        
        # Create temp PDF file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        # Export to PDF with mapping
        # Now supporting optional analysis_data passed to export method
        success = pdf_exporter.export(
            syllabus_data, 
            pdf_path, 
            include_mapping=True,
            analysis_data=analysis_data
        )
        
        if not success:
            raise HTTPException(status_code=500, detail="PDF export failed")
        
        # Return PDF file
        filename = f"{syllabus_data.get('course_code', 'syllabus')}_analysis.pdf" if analysis_data else f"{syllabus_data.get('course_code', 'syllabus')}.pdf"
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename
        )
        
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export/latex-pdf")
async def export_latex_pdf(request: OptimizeRequest, auth=Depends(verify_api_key)):
    """
    Export syllabus to PDF using LaTeX
    
    Features:
    - Professional academic formatting
    - Proper math formula rendering (O(n²), Big-O notation)
    - CO-PO mapping tables
    
    Returns: PDF file (or .tex if LaTeX not installed)
    """
    if not LATEX_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="LaTeX exporter not available. Install with: pip install pylatex"
        )
    
    try:
        # Extract syllabus data from request
        syllabus_data = request.syllabus_data
        
        # Handle nested structure - if data is wrapped, unwrap it
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        logger.info(f"Exporting LaTeX PDF for course: {syllabus_data.get('course_title', 'N/A')}")
        
        # Create temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        # Export to PDF using LaTeX
        latex_exporter = LaTeXExporter()
        result_path = latex_exporter.export_pdf(syllabus_data, pdf_path)
        
        # Determine media type and filename based on result
        if result_path.endswith('.tex'):
            media_type = "text/x-tex"
            filename = f"{syllabus_data.get('course_code', 'syllabus')}.tex"
        else:
            media_type = "application/pdf"
            filename = f"{syllabus_data.get('course_code', 'syllabus')}_latex.pdf"
        
        return FileResponse(
            result_path,
            media_type=media_type,
            filename=filename
        )
        
    except Exception as e:
        logger.error(f"LaTeX PDF export failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))



@app.post("/api/extract-outcomes")
async def extract_outcomes(text: str, auth=Depends(verify_api_key)):
    """
    Extract learning outcomes from text
    
    Returns: Extracted and validated outcomes
    """
    try:
        outcomes = outcome_extractor.extract_outcomes(text)
        
        return {
            "success": True,
            "outcomes": outcomes
        }
        
    except Exception as e:
        logger.error(f"Outcome extraction failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/validate-outcome")
async def validate_outcome(outcome: str, auth=Depends(verify_api_key)):
    """
    Validate a learning outcome
    
    Returns: Validation report with suggestions
    """
    try:
        validation = outcome_extractor.validate_outcome(outcome)
        
        return {
            "success": True,
            "validation": validation
        }
        
    except Exception as e:
        logger.error(f"Validation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export/excel")
async def export_excel(request: OptimizeRequest, auth=Depends(verify_api_key)):
    """Export syllabus to Excel"""
    try:
        syllabus_data = request.syllabus_data
        
        # Handle nested structure - if data is wrapped, unwrap it
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            # Frontend sent upload response structure
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            # Frontend sent optimize response structure
            syllabus_data = syllabus_data['syllabus']
        
        # Log what we received
        logger.info(f"Exporting Excel for course: {syllabus_data.get('course_title', 'N/A')}")
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            output_path = tmp.name
            
        excel_exporter = ExcelExporter()
        success = excel_exporter.export_complete_syllabus(
            syllabus_data, output_path, include_mapping=True, include_rubrics=True
        )
        
        if not success:
            logger.error(f"Excel export failed for '{syllabus_data.get('course_title', 'Unknown')}'")
            raise HTTPException(status_code=500, detail="Excel export failed")
            
        logger.info(f"Excel export successful for '{syllabus_data.get('course_title', 'Unknown')}'")
        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            filename=f"{syllabus_data.get('course_code', 'syllabus')}.xlsx"
        )
    except Exception as e:
        logger.error(f"Excel export failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/export/word")
async def export_word(request: OptimizeRequest, auth=Depends(verify_api_key)):
    """Export syllabus to Word document with academic formatting"""
    if not DOCX_AVAILABLE:
        raise HTTPException(
            status_code=501,
            detail="Word export not available. python-docx package is not installed."
        )
    
    try:
        syllabus_data = request.syllabus_data
        analysis_data = request.analysis_data  # Get optimization results
        
        # Handle nested structure
        if 'data' in syllabus_data and isinstance(syllabus_data.get('data'), dict):
            syllabus_data = syllabus_data['data']
        elif 'syllabus' in syllabus_data and isinstance(syllabus_data.get('syllabus'), dict):
            syllabus_data = syllabus_data['syllabus']
        
        # Merge analysis_data fields into syllabus_data if present
        # This ensures CO-PO mapping and other optimization results are available
        if analysis_data:
            if analysis_data.get('co_po_mapping') and not syllabus_data.get('co_po_mapping'):
                syllabus_data['co_po_mapping'] = analysis_data['co_po_mapping']
            if analysis_data.get('bloom_analysis') and not syllabus_data.get('bloom_analysis'):
                syllabus_data['bloom_analysis'] = analysis_data['bloom_analysis']
        
        logger.info(f"Exporting Word for course: {syllabus_data.get('course_title', 'N/A')}")
        
        # Create Word document
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            output_path = tmp.name
        
        doc = Document()
        
        # Import shared styles needed for tables/paragraphs
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

        # --- Institution Details Header ---
        if syllabus_data.get('university_name'):
            uni = doc.add_paragraph(syllabus_data.get('university_name'))
            uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = uni.runs[0]
            run.font.size = Pt(16)
            run.font.bold = True
            
        if syllabus_data.get('faculty_name'):
            fac = doc.add_paragraph(syllabus_data.get('faculty_name'))
            fac.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fac.runs[0]
            run.font.size = Pt(12)
            
        if syllabus_data.get('program') or syllabus_data.get('department'):
            prog_text = syllabus_data.get('program') or syllabus_data.get('department', '')
            prog = doc.add_paragraph(prog_text)
            prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = prog.runs[0]
            run.font.size = Pt(11)
            run.font.italic = True
            
        doc.add_paragraph()  # Spacer

        # --- Course Title ---
        course_code = syllabus_data.get('course_code', '')
        course_title = syllabus_data.get('course_title', 'Course Syllabus')
        title_text = f"{course_code}: {course_title}" if course_code else course_title
        
        title_para = doc.add_paragraph(title_text)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102) # Dark blue
        
        doc.add_paragraph() # Spacer
        
        # --- Course Info Table ---
        # Create a 2x2 table for course metadata
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Row 1: Course Type | Semester
        row0 = table.rows[0]
        row0.cells[0].text = f"Course Type: {syllabus_data.get('course_type', 'DSC')}"
        row0.cells[1].text = f"Semester: {syllabus_data.get('semester', 'I')}"
        
        # Row 2: Credits | Year
        credits = syllabus_data.get('credits', '3-0-0')
        year = syllabus_data.get('year', '')
        row1 = table.rows[1]
        row1.cells[0].text = f"Credits (L-T-P): {credits}"
        row1.cells[1].text = f"Academic Year: {year}" if year else ""
        
        # Bold the keys in the table cells
        for row in table.rows:
            for cell in row.cells:
                # Basic styling hook - iterate paragraphs to bold prefix if needed
                # (Simplified here as text assignment overwrites runs)
                p = cell.paragraphs[0]
                run = p.runs[0]
                run.font.size = Pt(10)
        
        doc.add_paragraph() # Spacer
        
        # --- Overview ---
        if syllabus_data.get('overview'):
            h = doc.add_heading('Course Overview', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            doc.add_paragraph(syllabus_data['overview'])
            doc.add_paragraph()

        # --- Learning Outcomes (Table Format) ---
        if syllabus_data.get('learning_outcomes'):
            h = doc.add_heading('Course Learning Outcomes (COs)', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Grid Table 4 Accent 1' # Built-in style
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'CO'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = "Bloom's Level"
            
            # Rows
            for outcome in syllabus_data['learning_outcomes']:
                row_cells = table.add_row().cells
                row_cells[0].text = outcome.get('code', '')
                row_cells[1].text = outcome.get('description', '')
                row_cells[2].text = outcome.get('bloom_level', '').capitalize()
            
            # Set widths
            table.autofit = False
            table.columns[0].width = Inches(0.8)
            table.columns[1].width = Inches(3.8)
            table.columns[2].width = Inches(1.4)
            
            doc.add_paragraph()

        # --- Units (Table Format) ---
        if syllabus_data.get('units'):
            h = doc.add_heading('Course Content', level=1)
            h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            for unit in syllabus_data['units']:
                # Create a 2-column table for Unit Header (Title | Hours)
                # We can simulate the "gray header" look by just using bold text or a table
                unit_num = unit.get('unit_number', '')
                title = unit.get('title', 'Untitled')
                hours = unit.get('hours', 0)
                
                # Unit Header Table (1 row, 2 cols)
                utable = doc.add_table(rows=1, cols=2)
                utable.allow_autofit = False
                utable.columns[0].width = Inches(5.0)
                utable.columns[1].width = Inches(1.0)
                
                # Cell 1: Unit Title
                c1 = utable.rows[0].cells[0]
                p = c1.paragraphs[0]
                run = p.add_run(f"Unit {unit_num}: {title}")
                run.bold = True
                run.font.size = Pt(11)
                
                # Cell 2: Hours
                c2 = utable.rows[0].cells[1]
                p = c2.paragraphs[0]
                run = p.add_run(f"{hours} Hours")
                run.bold = True
                run.font.size = Pt(11)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Shading (Background color) requires OXML/XML manipulation in python-docx
                # Keeping it simple with text formatting for now and a bottom border if possible
                
                # Topics List below the header
                topics = unit.get('topics', [])
                if topics:
                    topics_text = ", ".join(topics)
                    p = doc.add_paragraph(topics_text)
                    p.paragraph_format.left_indent = Inches(0.2)
                    p.paragraph_format.space_after = Pt(12)
                else:
                    doc.add_paragraph()

        # --- References ---
        if syllabus_data.get('references'):
             h = doc.add_heading('References', level=1)
             h.runs[0].font.color.rgb = RGBColor(44, 62, 80)
             
             refs = syllabus_data.get('references', {})
             
             if isinstance(refs, dict):
                 # Textbooks
                 if refs.get('textbooks'):
                     doc.add_heading('Textbooks:', level=2)
                     for t in refs['textbooks']:
                         doc.add_paragraph(t, style='List Bullet')
                 # Reference Books
                 if refs.get('references'):
                     doc.add_heading('Reference Books:', level=2)
                     for r in refs['references']:
                         doc.add_paragraph(r, style='List Bullet')
                 # Online Resources
                 if refs.get('online_resources'):
                     doc.add_heading('Online Resources:', level=2)
                     for r in refs['online_resources']:
                        doc.add_paragraph(r, style='List Bullet')
             elif isinstance(refs, list):
                 for r in refs:
                     doc.add_paragraph(r, style='List Bullet')

        # Save document
        doc.save(output_path)
        logger.info(f"Word export successful for '{syllabus_data.get('course_title', 'Unknown')}'")
        
        return FileResponse(
            output_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"{syllabus_data.get('course_code', 'syllabus')}.docx"
        )
        
    except Exception as e:
        logger.error(f"Word export failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
